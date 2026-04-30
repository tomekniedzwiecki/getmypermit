"""
Wersja 2: prostszy template bez problematycznych pętli table.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), '..', '_spike_test_template_v2.docx')

d = Document()
title = d.add_heading('Test docx-templates spike', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
p1 = d.add_paragraph()
p1.add_run('Klient: ').bold = True
p1.add_run('+++full_client_name+++')

p2 = d.add_paragraph()
p2.add_run('Numer sprawy: ').bold = True
p2.add_run('+++case_number+++')

p3 = d.add_paragraph()
p3.add_run('Data: ').bold = True
p3.add_run('+++today+++')

d.add_paragraph()
d.add_paragraph('Test polskich znaków: ąćęłńóśźż ĄĆĘŁŃÓŚŹŻ')

# Pętla: każdy wiersz osobny paragraph (nie tabela)
d.add_heading('Harmonogram płatności', 1)
d.add_paragraph('+++FOR i IN installments+++')
p = d.add_paragraph()
p.add_run('Rata #').bold = True
p.add_run('+++$i.number+++ — +++$i.amount+++ zł — termin: +++$i.due_date+++ — +++$i.status+++')
d.add_paragraph('+++END-FOR+++')

d.save(OUT)
print(f"Saved: {OUT}, size: {os.path.getsize(OUT)} bytes")
