"""
Tworzy 4 podstawowe szablony DOCX dla Etapu II-A:
1. karta_przyjecia.docx        - karta przyjęcia sprawy
2. harmonogram_platnosci.docx  - harmonogram rat
3. pelnomocnictwo_klient_pl.docx - pełnomocnictwo cudzoziemca
4. instrukcja_klient.docx       - instrukcja co przygotować

Stack: docx-templates delim '+++'.
Pamięć: pętle FOR-END-FOR przez tabelę python-docx mogą się dzielić na XML runs.
Workaround: pisać prosto, w jednej komórce / paragraph.

Po wygenerowaniu uruchom:
    node scripts/upload_base_templates.mjs
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs', 'templates')
os.makedirs(OUT_DIR, exist_ok=True)


# ============================================================================
# 1. karta_przyjecia.docx
# ============================================================================
d = Document()
title = d.add_heading('KARTA PRZYJĘCIA SPRAWY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
d.add_paragraph('Numer sprawy: +++case_number+++').runs[0].bold = True
d.add_paragraph('Data przyjęcia: +++today_pl+++')

d.add_paragraph()
d.add_heading('Dane cudzoziemca', 1)
d.add_paragraph('Imię i nazwisko: +++full_client_name+++')
d.add_paragraph('Data urodzenia: +++client_birth_date+++')
d.add_paragraph('Obywatelstwo: +++client_nationality+++')
d.add_paragraph('Telefon: +++client_phone+++')
d.add_paragraph('E-mail: +++client_email+++')

d.add_paragraph()
d.add_heading('Dane sprawy', 1)
d.add_paragraph('Kategoria: +++category_label+++')
d.add_paragraph('Tryb sprawy: +++kind_label+++')
d.add_paragraph('Pracodawca: +++employer_name+++')

d.add_paragraph()
d.add_heading('Opłaty planowane', 1)
d.add_paragraph('Wynagrodzenie kancelarii: +++fee_amount+++ zł')
d.add_paragraph('Opłata administracyjna: +++admin_fee_amount+++ zł')
d.add_paragraph('Opłata za kartę pobytu: +++stamp_fee_amount+++ zł')

d.add_paragraph()
d.add_paragraph().add_run('Podpis pracownika kancelarii: ___________________________').italic = True
d.add_paragraph().add_run('Podpis klienta: ___________________________').italic = True

d.save(os.path.join(OUT_DIR, 'karta_przyjecia.docx'))
print('✓ karta_przyjecia.docx')


# ============================================================================
# 2. harmonogram_platnosci.docx — bez FOR (pierwsze 6 rat hardcoded jako preview)
# ============================================================================
# UWAGA: docx-templates FOR loops w tabelach python-docx mają problem z XML runs.
# Workaround: w pierwszym deploycie obsługujemy tylko liczbę rat, łączną kwotę i daty
# pierwszej/ostatniej. Pełna tabela zostanie ręcznie sformatowana w Word lub
# wygenerowana przez Python wcześniej (alternatywne podejście dla pętli).
d = Document()
title = d.add_heading('HARMONOGRAM PŁATNOŚCI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
d.add_paragraph('Numer sprawy: +++case_number+++').runs[0].bold = True
d.add_paragraph('Klient: +++full_client_name+++')
d.add_paragraph('Data: +++today_pl+++')

d.add_paragraph()
d.add_heading('Podsumowanie', 1)
d.add_paragraph('Liczba rat: +++installments_count+++')
d.add_paragraph('Łączna kwota wynagrodzenia: +++installments_total+++ zł')

d.add_paragraph()
d.add_paragraph(
    'Szczegółowy harmonogram rat dostępny w systemie CRM '
    '(zakładka Finanse karty sprawy). W razie pytań prosimy o kontakt.'
)

d.add_paragraph()
d.add_paragraph().add_run(
    'Wpłaty prosimy realizować na konto kancelarii zgodnie z fakturą.'
).italic = True

d.save(os.path.join(OUT_DIR, 'harmonogram_platnosci.docx'))
print('✓ harmonogram_platnosci.docx')


# ============================================================================
# 3. pelnomocnictwo_klient_pl.docx
# ============================================================================
d = Document()
title = d.add_heading('PEŁNOMOCNICTWO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
loc = d.add_paragraph()
loc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
loc.add_run('+++today_pl+++')

d.add_paragraph()
d.add_paragraph(
    'Ja, niżej podpisany(-a) +++full_client_name+++, urodzony(-a) +++client_birth_date+++, '
    'obywatelstwa +++client_nationality+++, '
    'niniejszym udzielam pełnomocnictwa Kancelarii GetMyPermit '
    'do reprezentowania mnie w sprawie pobytowej +++case_number+++ '
    '(kategoria: +++category_label+++) przed właściwymi organami '
    'administracji publicznej Rzeczypospolitej Polskiej.'
)

d.add_paragraph()
d.add_paragraph(
    'Pełnomocnictwo obejmuje w szczególności:'
)
d.add_paragraph('— składanie wniosków, pism, dokumentów i oświadczeń;', style='List Bullet')
d.add_paragraph('— odbiór decyzji, postanowień i pism;', style='List Bullet')
d.add_paragraph('— zapoznanie z aktami sprawy;', style='List Bullet')
d.add_paragraph('— ustanawianie pełnomocnika substytucyjnego.', style='List Bullet')

d.add_paragraph()
d.add_paragraph('Pełnomocnictwo jest ważne do czasu zakończenia sprawy lub jego pisemnego cofnięcia.')

d.add_paragraph()
d.add_paragraph()
sig = d.add_paragraph()
sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig.add_run('___________________________').italic = True
sig2 = d.add_paragraph()
sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sig2.add_run('(podpis czytelny mocodawcy)').italic = True

d.save(os.path.join(OUT_DIR, 'pelnomocnictwo_klient_pl.docx'))
print('✓ pelnomocnictwo_klient_pl.docx')


# ============================================================================
# 4. instrukcja_klient.docx
# ============================================================================
d = Document()
title = d.add_heading('INSTRUKCJA DLA KLIENTA', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
d.add_paragraph('Numer sprawy: +++case_number+++').runs[0].bold = True
d.add_paragraph('Klient: +++full_client_name+++')
d.add_paragraph('Kategoria sprawy: +++category_label+++')
d.add_paragraph('Data: +++today_pl+++')

d.add_paragraph()
d.add_heading('Co należy przygotować', 1)
d.add_paragraph(
    'Szanowny Kliencie, w związku z prowadzeniem Twojej sprawy '
    'prosimy o przygotowanie następujących dokumentów. '
    'Szczegółowa lista wymaganych dokumentów dla kategorii sprawy +++category_label+++ '
    'jest dostępna w systemie CRM (zakładka Checklista karty sprawy).'
)

d.add_paragraph()
d.add_paragraph('Dokumenty podstawowe (zawsze wymagane):')
d.add_paragraph('— ważny paszport (oryginał + kopie wszystkich zapisanych stron);', style='List Bullet')
d.add_paragraph('— 4 aktualne fotografie biometryczne (3,5 × 4,5 cm);', style='List Bullet')
d.add_paragraph('— dowód uiszczenia opłaty administracyjnej (+++admin_fee_amount+++ zł);', style='List Bullet')
d.add_paragraph('— dowód uiszczenia opłaty za kartę pobytu (+++stamp_fee_amount+++ zł);', style='List Bullet')
d.add_paragraph('— podpisane pełnomocnictwo dla kancelarii.', style='List Bullet')

d.add_paragraph()
d.add_heading('Kontakt', 1)
d.add_paragraph(
    'W razie pytań prosimy o kontakt z opiekunem sprawy:'
)
d.add_paragraph('Twój e-mail: +++client_email+++')
d.add_paragraph('Twój telefon: +++client_phone+++')

d.add_paragraph()
d.add_paragraph().add_run(
    'Niniejsza instrukcja ma charakter informacyjny i nie zastępuje '
    'pełnej listy wymaganych dokumentów dla danej kategorii sprawy.'
).italic = True

d.save(os.path.join(OUT_DIR, 'instrukcja_klient.docx'))
print('✓ instrukcja_klient.docx')

print(f'\nWszystkie 4 szablony zapisane w: {OUT_DIR}')
print('\nNext step: node scripts/upload_base_templates.mjs')
