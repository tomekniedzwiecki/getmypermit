"""
Wersja 2 — bez FOR loop, używa pre-renderowanego pola checklists_as_text.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), '..', 'docs', 'templates', 'audit_checklist_universal.docx')

d = Document()
title = d.add_heading('AUDYT SPRAWY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
p = d.add_paragraph()
p.add_run('Numer sprawy: ').bold = True
p.add_run('+++case_number+++')

p = d.add_paragraph()
p.add_run('Klient: ').bold = True
p.add_run('+++full_client_name+++')

p = d.add_paragraph()
p.add_run('Kategoria: ').bold = True
p.add_run('+++category_label+++')

p = d.add_paragraph()
p.add_run('Tryb: ').bold = True
p.add_run('+++kind_label+++')

p = d.add_paragraph()
p.add_run('Pracodawca: ').bold = True
p.add_run('+++employer_name+++')

p = d.add_paragraph()
p.add_run('Data audytu: ').bold = True
p.add_run('+++today_pl+++')

d.add_paragraph()
d.add_paragraph(
    'Status pozycji: [V] zrobione  |  [ ] do zrobienia  |  [-] nie dotyczy  |  [!] zablokowane'
).italic = True

d.add_paragraph()
d.add_heading('Checklista pozycji', 1)

# Pre-renderowana lista jako jeden placeholder — działa zawsze
# (workaround: python-docx dzieli FOR loops na XML runs)
d.add_paragraph('+++checklists_as_text+++')

d.add_paragraph()
d.add_paragraph('---').alignment = WD_ALIGN_PARAGRAPH.CENTER

d.add_paragraph()
p = d.add_paragraph()
p.add_run('Audyt sporządzony przez Kancelarię GetMyPermit.').italic = True

d.save(OUT)
print(f'✓ Saved: {OUT}, {os.path.getsize(OUT)} bytes')
