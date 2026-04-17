"""
Preview struktury plikow od Pawla - pokazuje arkusze, tabele, kolumny i pierwsze wiersze.
Zapisuje wszystko LOKALNIE do folderu dane_od_pawla/_preview/ (gitignored).
"""
from pathlib import Path
import json
from datetime import datetime, timedelta
import numbers_parser.cell as _np_cell
from numbers_parser.cell import Cell, DateCell, EPOCH

# Monkey-patch _from_storage, bo w nowych plikach Numbers (26.x) niektore komorki daty
# maja wartosci poza zakresem datetime (OverflowError). Zwracamy tekst zamiast rzucac bledem.
_original_from_storage = Cell._from_storage

def _safe_from_storage(table_id, row, col, buffer, model):
    try:
        return _original_from_storage(table_id, row, col, buffer, model)
    except OverflowError:
        from numbers_parser.cell import TextCell
        return TextCell(row, col, "<date_overflow>")
    except Exception as e:
        from numbers_parser.cell import TextCell
        return TextCell(row, col, f"<err:{type(e).__name__}>")

Cell._from_storage = staticmethod(_safe_from_storage)

from numbers_parser import Document
from openpyxl import load_workbook
from docx import Document as DocxDocument

DANE = Path(__file__).parent.parent / "dane_od_pawla"
OUT = DANE / "_preview"
OUT.mkdir(exist_ok=True)

def _cell_value_safe(cell):
    """Bezpieczne odczytanie wartosci komorki - lapie bledy dat."""
    try:
        v = cell.value
        return str(v) if v is not None else ""
    except (OverflowError, ValueError, TypeError) as e:
        # Fallback - probuj wyciagnac raw
        try:
            return f"<err:{type(e).__name__}>"
        except Exception:
            return "<unreadable>"

def preview_numbers(path: Path):
    doc = Document(str(path))
    result = {"file": path.name, "type": "numbers", "sheets": []}
    try:
        sheets = list(doc.sheets)
    except Exception as e:
        return {**result, "error": f"doc.sheets failed: {type(e).__name__}: {e}"}
    for sheet in sheets:
        sheet_info = {"name": sheet.name, "tables": []}
        try:
            tables = list(sheet.tables)
        except Exception as e:
            sheet_info["error"] = f"sheet.tables failed: {type(e).__name__}: {e}"
            result["sheets"].append(sheet_info)
            continue
        for table in tables:
            # Czytamy komorkami zeby obejsc OverflowError na konkretnych datach
            all_rows = []
            try:
                num_rows = table.num_rows
                num_cols = table.num_cols
            except Exception:
                num_rows, num_cols = 0, 0
            for r in range(num_rows):
                row_values = []
                for c in range(num_cols):
                    try:
                        cell = table.cell(r, c)
                        row_values.append(_cell_value_safe(cell))
                    except Exception as e:
                        row_values.append(f"<cell_err:{type(e).__name__}>")
                all_rows.append(row_values)
            headers = all_rows[0] if all_rows else []
            sample = all_rows[1:6]
            sheet_info["tables"].append({
                "table_name": table.name,
                "rows": len(all_rows),
                "cols": len(headers),
                "headers": headers,
                "sample_first_5": sample,
            })
        result["sheets"].append(sheet_info)
    return result

def preview_xlsx(path: Path):
    wb = load_workbook(str(path), read_only=True, data_only=True)
    result = {"file": path.name, "type": "xlsx", "sheets": []}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        headers = [str(v) if v is not None else "" for v in (rows[0] if rows else [])]
        sample = [
            [str(v) if v is not None else "" for v in row]
            for row in rows[1:6]
        ]
        result["sheets"].append({
            "name": sheet,
            "rows": len(rows),
            "cols": len(headers),
            "headers": headers,
            "sample_first_5": sample,
        })
    return result

def preview_docx(path: Path):
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_preview = []
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows[:5]:
            rows_data.append([cell.text for cell in row.cells])
        tables_preview.append({
            "table_idx": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "sample_first_5": rows_data,
        })
    return {
        "file": path.name,
        "type": "docx",
        "paragraphs_count": len(paragraphs),
        "paragraphs_first_30": paragraphs[:30],
        "tables_count": len(doc.tables),
        "tables_preview": tables_preview,
    }

def main():
    for path in sorted(DANE.iterdir()):
        if path.is_dir() or path.name.startswith("_"):
            continue
        print(f"\n{'='*70}")
        print(f"PLIK: {path.name} ({path.stat().st_size // 1024} KB)")
        print('='*70)
        try:
            if path.suffix.lower() == ".numbers":
                data = preview_numbers(path)
            elif path.suffix.lower() == ".xlsx":
                data = preview_xlsx(path)
            elif path.suffix.lower() == ".docx":
                data = preview_docx(path)
            else:
                print(f"  (pomijam - nieznany format)")
                continue
            out_file = OUT / f"{path.stem}.preview.json"
            out_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            # Krotkie podsumowanie na stdout
            if data["type"] in ("numbers", "xlsx"):
                sheets = data["sheets"]
                for s in sheets:
                    if data["type"] == "numbers":
                        for t in s["tables"]:
                            print(f"  Arkusz '{s['name']}' / tabela '{t['table_name']}': {t['rows']} wierszy x {t['cols']} kolumn")
                            print(f"    Kolumny: {t['headers']}")
                    else:
                        print(f"  Arkusz '{s['name']}': {s['rows']} wierszy x {s['cols']} kolumn")
                        print(f"    Kolumny: {s['headers']}")
            elif data["type"] == "docx":
                print(f"  Paragrafy: {data['paragraphs_count']}, Tabele: {data['tables_count']}")
                for t in data["tables_preview"]:
                    print(f"  Tabela #{t['table_idx']}: {t['rows']} wierszy x {t['cols']} kolumn")
            print(f"  -> Preview zapisany: {out_file.relative_to(DANE.parent)}")
        except Exception as e:
            print(f"  BLAD: {type(e).__name__}: {e}")

if __name__ == "__main__":
    main()
