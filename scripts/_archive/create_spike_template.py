"""
Tworzy testowy template DOCX dla spike'u docx-templates.
Output: _spike_test_template.docx
Upload do bucket: document-templates (po jego utworzeniu w Etapie II-A)

Uruchomienie:
    python scripts/create_spike_template.py

Wymaga: pip install python-docx
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

import os

OUTPUT = os.path.join(os.path.dirname(__file__), '..', '_spike_test_template.docx')

d = Document()

# Tytuł
title = d.add_heading('Test docx-templates spike', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadane sprawy
d.add_paragraph()
p1 = d.add_paragraph()
p1.add_run('Klient: ').bold = True
p1.add_run('{full_client_name}')

p2 = d.add_paragraph()
p2.add_run('Numer sprawy: ').bold = True
p2.add_run('{case_number}')

p3 = d.add_paragraph()
p3.add_run('Data: ').bold = True
p3.add_run('{today}')

# Harmonogram rat
d.add_paragraph()
d.add_heading('Harmonogram płatności', 1)

table = d.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Lp.'
hdr[1].text = 'Kwota (zł)'
hdr[2].text = 'Termin'
hdr[3].text = 'Status'

# docx-templates: pętla — cała linia w komórkach
# {FOR i IN installments} ... {END-FOR}
row = table.add_row().cells
row[0].text = '{FOR i IN installments}{$i.number}'
row[1].text = '{$i.amount}'
row[2].text = '{$i.due_date}'
row[3].text = '{$i.status}{END-FOR}'

# Test polskich znaków
d.add_paragraph()
d.add_paragraph('Test polskich znaków: ąćęłńóśźż ĄĆĘŁŃÓŚŹŻ')

# Test conditional
d.add_paragraph()
d.add_paragraph('{IF client.first_name == "Jan"}Specjalna wiadomość dla Jana.{END-IF}')

d.save(OUTPUT)
print(f"Saved: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT)} bytes")
print()
print("Next step: upload do bucket 'document-templates' w Supabase Storage")
print("(Po utworzeniu bucketu w Etapie II-A migracja 20260507_05_storage_buckets.sql)")
