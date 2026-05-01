"""V3: TYLKO proste placeholdery, bez FOR/IF, żeby potwierdzić że render działa end-to-end."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

OUT = os.path.join(os.path.dirname(__file__), '..', '_spike_test_template_v3.docx')
d = Document()
d.add_heading('Spike DOCX v3 — proste placeholdery', 0)
d.add_paragraph('Klient: +++full_client_name+++')
d.add_paragraph('Numer sprawy: +++case_number+++')
d.add_paragraph('Data: +++today+++')
d.add_paragraph('Polskie znaki: ąćęłńóśźż ĄĆĘŁŃÓŚŹŻ')
d.save(OUT)
print(f"Saved: {OUT}, {os.path.getsize(OUT)} bytes")
